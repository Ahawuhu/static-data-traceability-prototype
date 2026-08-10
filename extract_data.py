"""Extract all meaningful table rows from the six Blue Island DOCX files."""
import json
import re
from pathlib import Path
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

SOURCE = Path('/Users/zhixuanzhang/Desktop/智慧城/静态数据溯源/2026年静态溯源数据收集结果/北部区域项目静态数据业务填报/蓝岛')
OUTPUT = Path(__file__).with_name('trace-data.js')
TEMPLATE_SOURCE = Path('/Users/zhixuanzhang/Desktop/智慧城/静态数据溯源/静态数据填写模板-新')
CUSTOMER_TEMPLATE = TEMPLATE_SOURCE / '客服类_新.xlsx'
OVERVIEW_TEMPLATE = TEMPLATE_SOURCE / '小区概况_新.xlsx'
FACILITY_TEMPLATE = TEMPLATE_SOURCE / '设施设备类_新 260810xlsx.xlsx'
CUSTOMER_HOUSE_TYPE_SOURCE = Path('/Users/zhixuanzhang/Desktop/智慧城/静态数据溯源/2026年静态溯源数据收集结果/北部区域项目静态数据业务填报/民航/用户上传_用户上传_已填写文档_民航公寓-客服类_1777014587925_0_e8zp_with_remarks.docx')

def clean(value):
    return re.sub(r'\s+', ' ', value or '').strip()

def meaningful(row):
    values = [v for v in row if v and v not in {'无', '/', '序号', '合计'}]
    return bool(values)

def category_for(name):
    for category in ['财务类', '环境类', '客服类', '行政类', '秩序类']:
        if f'《{category}-' in name:
            return category
    return '设施设备类'

result = {}
for path in sorted(SOURCE.glob('用户上传_*.docx')):
    category = category_for(path.name)
    doc = Document(path)
    notes_by_table = {}
    headings_by_table = {}
    current_table = -1
    pending_heading = None
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_Tbl):
            current_table += 1
            if pending_heading:
                headings_by_table[current_table] = pending_heading
                pending_heading = None
        elif isinstance(child, CT_P) and current_table >= 0:
            paragraph_text = clean(Paragraph(child, doc).text)
            heading_match = re.match(r'^\d+\.\d+\s+(.+)$', paragraph_text)
            if heading_match:
                pending_heading = heading_match.group(1).strip()
            if paragraph_text.startswith(('注：', '注:', '备注：', '备注:')):
                notes_by_table.setdefault(current_table, []).append(paragraph_text)
    toc_names = []
    for paragraph in doc.paragraphs:
        text = clean(paragraph.text)
        match = re.match(r'^\d+\.\d+\s+(.+?)(?:\s+\d+)?$', text)
        if match:
            name = match.group(1).strip()
            if name not in toc_names:
                toc_names.append(name)
    tables = []
    for index, table in enumerate(doc.tables):
        raw = [[clean(cell.text) for cell in row.cells] for row in table.rows]
        if not raw:
            continue
        # The first small table in every file is project metadata.
        if index == 0:
            title = '项目填报信息'
        elif index == 1:
            title = f'{category}基本情况汇总表'
        elif index in headings_by_table:
            title = headings_by_table[index]
        elif index - 2 < len(toc_names):
            title = toc_names[index - 2]
        else:
            title = f'数据表 {index + 1}'
        if title.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、')):
            title = title.split('、', 1)[1]

        header_rows = 1
        prefix_rows = []
        if title == '车库车位清单':
            header = [
                '序号', '车位号', '车位性质-产权', '车位性质-人防', '房号',
                '产权人/租赁人姓名', '联系电话', '购买/租赁时间',
                '车位使用情况-业主购买产权-自用', '车位使用情况-业主购买产权-出租',
                '车位使用情况-业主购买产权-空置', '车位使用情况-房开出租',
                '车位使用情况-未出售（租）', '车牌号'
            ]
            header_rows = 3
        elif len(raw) > 1 and len(set(raw[0])) == 1:
            # Some detail sheets start with a full-width building/project heading.
            # Preserve that row as a section heading and use the next row as columns.
            header = raw[1]
            header_rows = 2
            prefix_rows = [raw[0]]
        else:
            header = raw[0]

        header = [value or f'字段{i + 1}' for i, value in enumerate(header)]
        header_corrections = {
            '消防通讯设备': {10: '维保/管理单位'},
            '图形显示装置（CRT）': {10: '维保/管理单位'},
            '消控室配套设备': {10: '维保/管理单位'},
            '地下车库通风系统': {10: '维保/管理单位'},
            '地下车库交通设施': {10: '维保/管理单位'},
        }
        for position, corrected_name in header_corrections.get(title, {}).items():
            if position < len(header):
                header[position] = corrected_name

        width = len(header)
        candidate_rows = prefix_rows + raw[header_rows:]
        rows = []
        for row in candidate_rows:
            normalized = (row + [''] * width)[:width]
            if normalized == header or not meaningful(normalized):
                continue
            rows.append(normalized)
        tables.append({'name': title[:36], 'columns': header, 'rows': rows, 'notes': notes_by_table.get(index, [])})
    result[category] = {
        'file': path.name,
        'tables': tables,
        'rowCount': sum(len(table['rows']) for table in tables)
    }

def template_schemas(path):
    """Return the authoritative sheet names and headers from a template workbook."""
    workbook = load_workbook(path, data_only=True, read_only=True)
    schemas = []
    for sheet in workbook.worksheets:
        values = [[clean(str(value)) if value is not None else '' for value in row]
                  for row in sheet.iter_rows(values_only=True)]
        if sheet.title == '项目基本信息':
            schemas.append(('项目填报信息', None))
            continue
        header_index = 2 if '小区户型表' in sheet.title else 1
        title = re.sub(r'^(?:\d+(?:\.\d+)?\s*|[一二三四五六七八九十百]+、)', '', sheet.title).strip()
        schemas.append((title, values[header_index]))
    return schemas

# The original customer DOCX contains the community overview as its second table.
# Split it into its own first category while preserving Blue Island's filled values.
customer = result['客服类']
original_customer_tables = customer['tables']
overview_schemas = template_schemas(OVERVIEW_TEMPLATE)
overview_tables = [original_customer_tables[0], original_customer_tables[1]]
overview_tables[0]['name'] = overview_schemas[0][0]
overview_tables[1]['name'] = overview_schemas[1][0]
overview_tables[1]['columns'] = overview_schemas[1][1]

# The remaining customer sheets use the new customer workbook as the field authority.
customer_tables = [original_customer_tables[0]] + original_customer_tables[2:]
customer_schemas = template_schemas(CUSTOMER_TEMPLATE)
if len(customer_tables) != len(customer_schemas):
    raise ValueError(f'客服类表数与新模板不一致: {len(customer_tables)} != {len(customer_schemas)}')
for table, (name, columns) in zip(customer_tables, customer_schemas):
    table['name'] = name
    if columns:
        if any(len(row) != len(columns) for row in table['rows']):
            raise ValueError(f'客服类表「{name}」数据列数与新模板不一致')
        table['columns'] = columns

# Use the completed Minhang Apartment house-type schedule for this one sheet.
house_type_document = Document(CUSTOMER_HOUSE_TYPE_SOURCE)
house_type_raw = [[clean(cell.text) for cell in row.cells] for row in house_type_document.tables[13].rows]
house_type_header = house_type_raw[1]
house_type_rows = []
for row in [house_type_raw[0]] + house_type_raw[2:]:
    normalized = (row + [''] * len(house_type_header))[:len(house_type_header)]
    if normalized == house_type_header or not meaningful(normalized):
        continue
    house_type_rows.append(normalized)
house_type_table = next(table for table in customer_tables if table['name'] == '小区户型表')
house_type_table['columns'] = house_type_header
house_type_table['rows'] = house_type_rows
house_type_table['notes'] = []
house_type_table['sourceFile'] = CUSTOMER_HOUSE_TYPE_SOURCE.name

customer['file'] = CUSTOMER_TEMPLATE.name
customer['tables'] = customer_tables
customer['rowCount'] = sum(len(table['rows']) for table in customer_tables)

# The facility workbook is the authoritative field definition for all facility tables.
# Keep the existing filled values where positions match, while normalizing every row
# to the template width so the rendered table cannot drift from its header.
facility = result['设施设备类']
facility_schemas = template_schemas(FACILITY_TEMPLATE)
facility_tables = facility['tables']
if len(facility_tables) != len(facility_schemas):
    raise ValueError(f'设施设备类表数与新模板不一致: {len(facility_tables)} != {len(facility_schemas)}')
for table, (name, columns) in zip(facility_tables[1:], facility_schemas[1:]):
    table['name'] = name
    if columns:
        width = len(columns)
        table['columns'] = columns
        table['rows'] = [(row + [''] * width)[:width] for row in table['rows']]
facility['file'] = FACILITY_TEMPLATE.name
facility['rowCount'] = sum(len(table['rows']) for table in facility_tables)
overview = {
    'file': OVERVIEW_TEMPLATE.name,
    'tables': overview_tables,
    'rowCount': sum(len(table['rows']) for table in overview_tables),
}
result = {'小区概况': overview, **result}

OUTPUT.write_text('window.TRACE_DATA=' + json.dumps(result, ensure_ascii=False, separators=(',', ':')) + ';\n', encoding='utf-8')
print(json.dumps({key: {'tables': len(value['tables']), 'rows': value['rowCount']} for key, value in result.items()}, ensure_ascii=False, indent=2))
